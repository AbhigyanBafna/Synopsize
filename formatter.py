import docx
from docx.oxml.ns import qn
from tkinter import filedialog
import datetime

def wordDoc(dayOfMonth, dayOfWeek, time, month, title, summary):

    # Get the current date and time
    now = datetime.datetime.now()

    # Define the variables
    if str(dayOfMonth) == "Day of Month":
        day_of_month = now.strftime("%d")

        suffix = ""
        if day_of_month in ["01", "21", "31"]:
            suffix = "st"
        elif day_of_month in ["02", "22"]:
            suffix = "nd"
        elif day_of_month in ["03", "23"]:
            suffix = "rd"
        else:
            suffix = "th"

        day_of_month = day_of_month + suffix

    else:
        day_of_month = str(dayOfMonth)

    if str(time) == "Time":
        time = now.strftime("%H:%M") + " hrs"
    else:
        time = str(time)

    if str(dayOfWeek) == "Day":
        day_of_week = now.strftime("%A")
    else:
        day_of_week = str(dayOfWeek)
    
    if str(month) == "Month":
        month = now.strftime("%B")
    else:
        month = str(month)

    if str(title) == "":
        title = "Summary"
    else:
        title = str(title)

    if str(summary) == "":
        summary = "Sorry, something went wrong :( Try again or report to Developer."
    else:
        summary = str(summary)


    year = now.year

    # Test Summary : "This is a summary of my document. It contains important information that I want to convey to the reader. Sarah is a sentence web developer. Altaf is a math genius. Amisha is Ausha. Abhigyan only builds cool stuff."

    # Create a new Word document
    document = docx.Document()

    # Add the date and time to the header
    header = document.sections[0].header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_paragraph.add_run()
    header_run.add_text(f"{day_of_month} {month} {year}").bold = True
    header_run.add_text("\t\t\t\t\t\t\t")
    header_run.add_text(f"{day_of_week}, {time}").bold = True

    # Add the title to the document
    title_paragraph = document.add_paragraph(title)
    title_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.style = "Title"


    # Add the summary to the document
    summary_paragraph = document.add_paragraph()
    sentences = summary.split(".")
    for sentence in sentences:
        if sentence:
            summary_bullet_points = summary_paragraph.add_run(f"\n• {sentence.strip()}.")
            summary_bullet_points.font.name = "Symbol"
            summary_bullet_points._element.rPr.rFonts.set(qn('w:eastAsia'), 'Symbol')
            summary_bullet_points._element.rPr.rFonts.set(qn('w:cs'), 'Symbol')


    # Ask the user for the file path to save the document
    filename = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=(("Word Document", "*.docx"), ("All files", "*.*")))

    # Save the document to the specified file path
    document.save(filename)

